import os
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import Counter, defaultdict

# Create document
doc = Document()
title = doc.add_heading('SQE 英国律师资格考试 - 必背考点汇总', 0)
subtitle = doc.add_paragraph('按科目整理 · 中英文对照 · 带解析')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

sqe_dir = r'C:\Users\Administrator\Desktop\Desktop\SQE'

# ==================== SQE必背考点.md content ====================
doc.add_heading('【核心考点】SQE必背考点', level=1)

key_points = """
## 合同法 CONTRACT LAW

### 有效合同5大要素
1. 要约 Offer - 向特定人发出的明确交易意愿
2. 承诺 Acceptance - 完全接受要约条件
3. 对价 Consideration - 有价值的交换
4. 缔约能力 Capacity - 双方有签订合同的法律能力
5. 合法性 Legality - 合同目的合法

### 要约与要约邀请的区别
- 要约 Offer：向特定人发出，收到后一旦承诺即成立合同
- 邀请要约 Invitation to Treat：邀请对方报价，如广告、橱窗展示、招标

### 对价3个规则
1. 对价必须存在（必须是有价值的）
2. 对价必须来自承诺方（必须是当事人自己付出）
3. 过去对价不算（Past consideration is no consideration）

### 合同无效情形
1. 错误 Misrepresentation - 重大事实错误
2. 胁迫 Duress - 强迫签订
3. 不当影响 Undue Influence - 滥用信任关系
4. 非法 Illegal - 目的违法
5. 虚假失信 Misconduct - 严重违约

### 赔偿金条款有效规则
- 违约金条款 Liquidated Damages：预先约定违约金额，如合理估算实际损失，则有效
- 罚金条款 Penalty：金额远超实际损失，则无效

### 侵权法 TORT LAW

### 过失侵权3要素
1. 注意义务 Duty of Care - 对原告负有合理注意义务
2. 违反义务 Breach - 未达到合理人标准
3. 因果关系 Causation - 违反义务导致损害

### 特殊侵权类型
- 疏忽 Negligence：注意义务+违反+损害+因果
- 滋扰 Nuisance：干扰土地使用或享受
- 非法侵入 Trespass：非法进入土地或干扰人身
- 名誉侵权 Defamation：诽谤或诬蔑导致名誉受损

### 雇主责任
是的，只要员工在受雇范围内（in the course of employment）行事，雇主承担替代责任Vicarious Liability

### 财产法 PROPERTY LAW

### 租赁权5大要素
1. 租户 Tenant - 明确
2. 房东 Landlord - 明确
3. 租期 Term - 书面（超过3年）
4. 租金 Rent - 如有需明确
5. 占有权 Possession - 转让使用权

### 信托3要素
1. 确定意图 Intention - 设立信托的明确意图
2. 确定财产 Certainty of subject matter - 信托财产明确
3. 确定受益人 Certainty of objects - 受益人明确或可确定

### 土地权益优先权（土地登记）
1. 法律权益 Legal Interest - 最优先
2. 衡平权益 Equitable Interest - 次优先
3. 登记顺序 - 先登记者优先

### 刑法 CRIMINAL LAW

### 犯罪构成要件
1. 行为 Actus Reus - 客观外在行为
2. 意图 Mens Rea - 主观故意或过失
3. 因果关系 Causation - 行为与结果有因果联系

### 谋杀 vs 过失杀人
- 谋杀 Murder：故意造成死亡或严重伤害（GBH with intent）
- 过失杀人 Manslaughter：非故意但造成死亡（可因疏忽或自愿导致）

### 辩护理由
1. 正当防卫 Self-defence - 合理自卫
2. 精神失常 Insanity - 精神疾病导致无法理解行为
3. 自愿醉酒 Voluntary intoxication - 一般不成立辩护
4. 胁迫 Duress - 被迫犯罪
5. 减轻情节 Mitigation - 情有可原

### 盗窃罪要素
1. 永久剥夺意图 - 永久性拿走他人财产
2. 偷窃意图 - 无权占有
3. 他人财产 - 财产属于他人
4. 实际拿走 - 实际控制他人财产

### 民事诉讼法 CIVIL PROCEDURE

### 诉讼时效
- 合同纠纷：6年
- 侵权纠纷：3年（人身伤害）
- 财产纠纷：12年

### 案件管理
- Small Claims Track - 小额诉讼（≤£10,000）
- Fast Track - 快速诉讼（£10,000-£25,000）
- Multi Track - 多轨诉讼（>£25,000 或复杂案件）

### 证明标准
民事案件：平衡可能性（on the balance of probabilities）

### 商业法 BUSINESS LAW

### 公司设立要件
1. 公司名称 - 唯一且不违法
2. 注册地址 - 英国注册地址
3. 公司章程 - 规定公司治理结构
4. 董事 - 至少一名董事
5. 股东 - 至少一名股东

### 董事职责
1. 诚信义务 - 为公司最佳利益行事
2. 注意义务 - 具备合理技能和谨慎
3. 忠实义务 - 避免利益冲突
4. 不越权 - 在公司章程范围内行事

### 破产程序
1. 自愿安排 IVA - 债务人与债权人协商
2. 清算 Liquidation - 公司资产变卖
3. 破产 Bankruptcy - 个人债务清偿程序
4. 管理 Administration - 公司重组程序

### 职业伦理 PROFESSIONAL CONDUCT

### 律师6大核心义务
1. 诚信 Honesty - 必须诚实守信
2. 独立 Independence - 保持独立，不受不当影响
3. 保密 Confidentiality - 严格保密客户信息
4. 利益冲突 Conflict of Interest - 避免利益冲突
5. 能力 Competence - 提供合格专业服务
6. 客户利益 Client's Best Interests - 以客户最佳利益行事

### 保密例外情况
1. 客户同意 - 客户明确同意
2. 法律要求 - 法院命令或法律规定
3. 防止犯罪 - 防止严重犯罪
4. 财务申报 - 财务情报要求

### 洗钱防范义务
1. 客户尽职调查 CDD - 核实客户身份
2. 可疑交易报告 STR - 报告可疑活动
3. 记录保存 - 保存交易记录至少5年
4. 员工培训 - 培训员工识别洗钱迹象
"""

for line in key_points.strip().split('\n'):
    if line.startswith('## '):
        doc.add_heading(line.replace('## ', ''), level=2)
    elif line.startswith('### '):
        doc.add_heading(line.replace('### ', ''), level=2)
    elif line.startswith('1. ') or line.startswith('2. ') or line.startswith('3. ') or line.startswith('4. ') or line.startswith('5. '):
        p = doc.add_paragraph(line)
    elif line.strip():
        p = doc.add_paragraph(line)

# ==================== From extracted questions ====================
doc.add_page_break()
doc.add_heading('【试题考点提取】按科目分类', level=1)

# Load extracted questions
eq_path = os.path.join(sqe_dir, 'extracted_questions_final.json')
with open(eq_path, 'r', encoding='utf-8') as f:
    eq_questions = json.load(f)

# Extract key points from questions
print("Analyzing questions for key points...")

# Group by FLK
flk1_questions = [q for q in eq_questions if q.get('cate') == 'FLK1']
flk2_questions = [q for q in eq_questions if q.get('cate') == 'FLK2']

print(f"FLK1: {len(flk1_questions)} questions")
print(f"FLK2: {len(flk2_questions)} questions")

# FLK1 Subjects
doc.add_heading('FLK1 科目考点', level=2)

flk1_subjects = {
    'Constitutional Law': '宪法与行政法',
    'Contract Law': '合同法',
    'Tort Law': '侵权法',
    'Business Law': '商法',
    'Criminal Law': '刑法',
    'EU Law': '欧盟法'
}

for subj_en, subj_cn in flk1_subjects.items():
    # Find related questions
    related = [q for q in flk1_questions if subj_en.lower() in str(q.get('q_en', '')).lower() or subj_en.lower() in str(q.get('q_cn', '')).lower()]
    if len(related) > 0:
        doc.add_heading(f'{subj_cn} ({subj_en}) - {len(related)}题', level=3)

        # Extract key legal points
        key_points_seen = set()
        for q in related[:100]:
            q_en = q.get('q_en', '')
            # Extract first meaningful question text
            lines = [l for l in q_en.split('\n') if l.strip() and len(l) > 30][:1]
            if lines:
                q_text = lines[0].strip()
                # Truncate for display
                if len(q_text) > 150:
                    q_text = q_text[:150] + '...'
                p = doc.add_paragraph()
                p.add_run('• ' + q_text)
        doc.add_paragraph()

# FLK2 Subjects
doc.add_heading('FLK2 科目考点', level=2)

flk2_subjects = {
    'Land Law': '土地法',
    'Trust Law': '信托法',
    'Wills and Probate': '遗嘱与遗产',
    'Civil Litigation': '民事诉讼',
    'Solicitors Accounts': '律师会计',
    'Legal Practice': '法律实务',
    'Ethics and Regulation': '伦理与监管'
}

for subj_en, subj_cn in flk2_subjects.items():
    # Find related questions
    related = [q for q in flk2_questions if subj_en.lower() in str(q.get('q_en', '')).lower() or subj_en.lower() in str(q.get('q_cn', '')).lower()]
    if len(related) > 0:
        doc.add_heading(f'{subj_cn} ({subj_en}) - {len(related)}题', level=3)

        for q in related[:100]:
            q_en = q.get('q_en', '')
            lines = [l for l in q_en.split('\n') if l.strip() and len(l) > 30][:1]
            if lines:
                q_text = lines[0].strip()
                if len(q_text) > 150:
                    q_text = q_text[:150] + '...'
                p = doc.add_paragraph()
                p.add_run('• ' + q_text)
        doc.add_paragraph()

# Save
output_path = r'C:\Users\Administrator\Desktop\SQE_必背考点_完整版.docx'
doc.save(output_path)
print(f"\nDocument saved to: {output_path}")