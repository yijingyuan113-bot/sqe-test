import os
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from collections import defaultdict

# Create document
doc = Document()
title = doc.add_heading('SQE 英国律师资格考试题目汇总', 0)
subtitle = doc.add_paragraph('中英文对照 · 中文解析 · 按科目整理')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

sqe_dir = r'C:\Users\Administrator\Desktop\Desktop\SQE'

# ==================== Source 1: question_bank/questions.json ====================
print("Loading question_bank/questions.json...")
qb_path = os.path.join(sqe_dir, 'question_bank', 'questions.json')
with open(qb_path, 'r', encoding='utf-8') as f:
    qb_questions = json.load(f)

# Group by category
qb_by_cat = defaultdict(list)
for q in qb_questions:
    cat = q.get('cate', 'Unknown')
    qb_by_cat[cat].append(q)

print(f"Question bank: {len(qb_questions)} questions")

# ==================== Source 2: extracted_questions_final.json ====================
print("Loading extracted_questions_final.json...")
eq_path = os.path.join(sqe_dir, 'extracted_questions_final.json')
with open(eq_path, 'r', encoding='utf-8') as f:
    eq_questions = json.load(f)

# Group by category
eq_by_cat = defaultdict(list)
for q in eq_questions:
    cat = q.get('cate', 'Unknown')
    eq_by_cat[cat].append(q)

print(f"Extracted questions: {len(eq_questions)} questions")

# ==================== Source 3: docx files ====================
print("Loading docx files...")
docx_questions = []

from docx import Document as DocxDoc

docx_files = [
    'FLK2真题（有答案）.docx',
    'Legal Practice Exam Questions.docx',
    'MCT Mock Test 2 模拟测试题解析.docx',
    'Practice Part A1.docx',
    'Practice Part B1.docx',
    'SQE Prep & Practise - Assignment有答案.docx',
    'SRA FLK1真题（有答案）.docx',
    'Simulated Exam Part 2 - FLK1.docx',
    'Simulated Exam Part 4 - FLK1.docx',
    'SQE.docx'
]

def extract_from_docx(filepath):
    try:
        d = DocxDoc(filepath)
        text = '\n'.join([p.text for p in d.paragraphs if p.text.strip()])
        return text
    except:
        return ""

# ==================== Compile by Subject ====================
print("\nOrganizing by subject...")

# Define subject mappings
subject_mapping = {
    'FLK1': {
        'Constitutional Law': '宪法',
        'Contract Law': '合同法',
        'Tort Law': '侵权法',
        'Business Law': '商法',
        'Criminal Law': '刑法',
        'EU Law': '欧盟法',
        'Legal Method': '法律方法',
        'Professional Conduct': '职业行为准则',
        'Civil Litigation': '民事诉讼',
        'Evidence': '证据法'
    },
    'FLK2': {
        'Land Law': '土地法',
        'Trust Law': '信托法',
        'Wills and Probate': '遗嘱与遗产',
        'Civil Litigation': '民事诉讼',
        'Solicitors Accounts': '律师会计',
        'Legal Practice': '法律实务',
        'Ethics and Regulation': '伦理与监管'
    }
}

# Process question bank (already has good structure)
print("\n=== Adding from question_bank ===")
for cat, questions in qb_by_cat.items():
    flk = 'FLK1' if 'FLK1' in str(cat) or cat == 'FLK1' else 'FLK2'
    subj = cat if cat in ['FLK1', 'FLK2'] else f"{flk} - {cat}"

    doc.add_heading(f'【{cat}】{len(questions)}题', level=1)

    for i, q in enumerate(questions[:200]):  # Limit 200 per category
        # Question text
        q_en = q.get('q_en', '')
        q_cn = q.get('q_cn', '')

        # Clean up question text
        if 'Question' in q_en:
            parts = q_en.split('Question')
            if len(parts) > 1:
                q_en = 'Question'.join(parts[1:])

        p = doc.add_paragraph()
        p.add_run(f"Q{i+1}: ").bold = True

        # Extract English question
        lines = q_en.split('\n')
        for line in lines[:3]:
            if line.strip():
                p.add_run(line.strip()[:200])
                break

        # Chinese translation
        if q_cn and len(q_cn) > 5:
            p = doc.add_paragraph()
            p.add_run(f"  中文: {q_cn[:300]}").font.color.rgb = RGBColor(0, 0, 255)

        # Options
        opts = q.get('options', [])
        for j, opt in enumerate(opts[:5]):
            if isinstance(opt, dict):
                opt_en = opt.get('en', opt.get('en_us', ''))
                opt_cn = opt.get('cn', '')
            else:
                opt_en = str(opt)
                opt_cn = ''

            p = doc.add_paragraph()
            opt_letter = chr(65 + j)  # A, B, C, D, E
            p.add_run(f"  {opt_letter}. {opt_en[:150]}")

            if opt_cn:
                p.add_run(f"\n     中文: {opt_cn[:100]}").font.color.rgb = RGBColor(0, 100, 0)

        # Answer
        ans = q.get('ans', '')
        if ans:
            p = doc.add_paragraph()
            p.add_run(f"  答案: {ans}").bold = True
            p.runs[-1].font.color.rgb = RGBColor(255, 0, 0)

        # Explanation
        exp_cn = q.get('exp_cn', '')
        exp_en = q.get('exp_en', '')
        if exp_cn:
            p = doc.add_paragraph()
            p.add_run(f"  解析: {exp_cn[:200]}").font.color.rgb = RGBColor(128, 0, 128)

        doc.add_paragraph('_' * 50)

# ==================== Process extracted_questions ====================
print("\n=== Adding from extracted_questions ===")

# Get categories not covered by question_bank
for cat, questions in eq_by_cat.items():
    if len(qb_by_cat.get(cat, [])) > 0:
        continue  # Skip if already covered

    if len(questions) > 500:
        questions = questions[:500]  # Limit large sets

    doc.add_heading(f'【{cat}】{len(questions)}题 (来源: extracted)', level=1)

    for i, q in enumerate(questions[:100]):
        q_en = q.get('q_en', '')
        q_cn = q.get('q_cn', '')

        p = doc.add_paragraph()
        p.add_run(f"Q{i+1}: ").bold = True

        # Get first meaningful line
        lines = [l for l in q_en.split('\n') if l.strip() and 'Question' not in l][:1]
        if lines:
            p.add_run(lines[0].strip()[:200])

        if q_cn and len(q_cn) > 5:
            p = doc.add_paragraph()
            p.add_run(f"  中文: {q_cn[:300]}").font.color.rgb = RGBColor(0, 0, 255)

        # Options
        opts = q.get('options', [])
        if isinstance(opts, list):
            for j, opt in enumerate(opts[:5]):
                if isinstance(opt, dict):
                    opt_en = opt.get('en', '')
                else:
                    opt_en = str(opt)
                p = doc.add_paragraph()
                opt_letter = chr(65 + j)
                p.add_run(f"  {opt_letter}. {opt_en[:150]}")

        doc.add_paragraph('_' * 50)

# Save
output_path = r'C:\Users\Administrator\Desktop\SQE_Complete_Organized.docx'
doc.save(output_path)
print(f"\nDocument saved to: {output_path}")

# Summary
print("\n=== Summary ===")
print(f"question_bank: {len(qb_questions)} questions")
print(f"extracted_questions: {len(eq_questions)} questions")
print("Total categories covered:", len(qb_by_cat) + len(eq_by_cat))