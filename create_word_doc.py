import json
import re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Read the questions file as text
with open('sqe_questions.js', 'r', encoding='utf-8-sig') as f:
    content = f.read()

# Remove the const declaration
content = content.replace('const sqeQuestions = ', '')

# Find all question objects using regex
# Each question starts with { and ends with }
questions = []
current_pos = 0
depth = 0
start = -1

for i, char in enumerate(content):
    if char == '{':
        if depth == 0:
            start = i
        depth += 1
    elif char == '}':
        depth -= 1
        if depth == 0 and start != -1:
            json_str = content[start:i+1]
            try:
                q = json.loads(json_str)
                questions.append(q)
            except:
                pass
            start = -1

print(f"Total questions found: {len(questions)}")

# Create Word document
doc = Document()
title = doc.add_heading('SQE FLK1 题目中英文对照及解析', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

flk1_count = 0
for q in questions:
    if q.get('flk') != 'FLK1':
        continue
    flk1_count += 1

    doc.add_heading(f'第{q.get("number", "?")}题 - {q.get("subject", "Unknown")}', level=1)

    # English stem
    p = doc.add_paragraph()
    run = p.add_run('英文原文: ')
    run.bold = True
    p.add_run(q.get('stem', ''))

    # Chinese stem
    stem_zh = q.get('stem_zh', '')
    if stem_zh:
        p = doc.add_paragraph()
        run = p.add_run('中文翻译: ')
        run.bold = True
        run.font.color.rgb = RGBColor(0, 0, 255)
        p.add_run(stem_zh).font.color.rgb = RGBColor(0, 0, 255)

    # Options
    opts = q.get('opts', {})
    for opt_key in ['A', 'B', 'C', 'D', 'E']:
        if opt_key in opts:
            opt_zh_key = f'{opt_key}_zh'
            doc.add_paragraph()
            p = doc.add_paragraph()
            run = p.add_run(f'{opt_key}. {opts[opt_key]}')
            run.bold = True

            if opt_zh_key in q:
                p = doc.add_paragraph()
                run = p.add_run(f'   中文: {q[opt_zh_key]}')
                run.font.color.rgb = RGBColor(0, 100, 0)

            if q.get('answer') == opt_key:
                p = doc.add_paragraph()
                run = p.add_run(f'   ✓ 正确答案')
                run.bold = True
                run.font.color.rgb = RGBColor(255, 0, 0)

    # Analysis
    analysis_zh = q.get('analysis_zh', '')
    if analysis_zh:
        doc.add_paragraph()
        p = doc.add_paragraph()
        run = p.add_run('解析: ')
        run.bold = True
        run.font.color.rgb = RGBColor(128, 0, 128)
        p.add_run(analysis_zh).font.color.rgb = RGBColor(128, 0, 128)

    doc.add_page_break()

# Save to Desktop
import os
desktop = os.path.join(os.path.expanduser('~'), 'Desktop', 'SQE_FLK1_题目解析.docx')
doc.save(desktop)
print(f'Document saved to: {desktop}')
print(f'Total FLK1 questions processed: {flk1_count}')