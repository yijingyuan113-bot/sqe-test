import os
import re
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create combined Word document
doc = Document()
title = doc.add_heading('SQE 英国律师资格考试题目汇总', 0)
subtitle = doc.add_paragraph('中英文对照 · 中文解析 · 按科目整理')
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()

# ================== Part 1: Desktop SQE folder ==================
doc.add_heading('第一部分：Desktop SQE 文件夹题目', level=1)

sqe_dir = r'C:\Users\Administrator\Desktop\Desktop\SQE'
docx_files = [f for f in os.listdir(sqe_dir) if f.endswith('.docx')
              and not f.startswith('words') and not f.startswith('千')
              and 'FLK' in f or 'SRA' in f or 'Practice' in f or 'Legal' in f]

def read_docx_full(path):
    doc = Document(path)
    return '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])

def parse_docx_questions(text):
    questions = []
    lines = text.split('\n')
    current_q = None
    section = 'start'

    for line in lines:
        line = line.strip()
        if not line:
            continue

        # Question number
        q_match = re.match(r'^(?:Question\s*)?(\d+)\s*$', line, re.IGNORECASE)
        if q_match:
            if current_q and current_q.get('english'):
                questions.append(current_q)
            current_q = {'number': int(q_match.group(1)), 'english': '', 'chinese': '', 'options': {}, 'answer': ''}
            section = 'question'
            continue

        if current_q:
            # Options
            opt_match = re.match(r'^([A-E])[.\)]\s*(.+)', line, re.IGNORECASE)
            if opt_match:
                current_q['options'][opt_match.group(1).upper()] = {
                    'english': opt_match.group(2).strip(),
                    'chinese': ''
                }
                continue

            # Answer
            if 'answer' in line.lower() or '正确答案' in line:
                ans_match = re.search(r'\b([A-E])\b', line, re.IGNORECASE)
                if ans_match:
                    current_q['answer'] = ans_match.group(1).upper()
                continue

            # Chinese line indicator
            if re.match(r'^[一-鿿]', line) and not current_q.get('chinese'):
                current_q['chinese'] = line
                continue

            # English content
            if section == 'question' or not current_q.get('chinese'):
                current_q['english'] += ' ' + line

    if current_q and current_q.get('english'):
        questions.append(current_q)
    return questions

for filename in docx_files:
    filepath = os.path.join(sqe_dir, filename)
    try:
        print(f"Processing: {filename}")
        text = read_docx_full(filepath)
        questions = parse_docx_questions(text)
        if questions:
            doc.add_heading(f'【{filename.replace(".docx","")}】共{len(questions)}题', level=2)
            for q in questions[:100]:
                p = doc.add_paragraph()
                p.add_run(f"Q{q['number']}: ").bold = True
                p.add_run(q.get('english', '')[:200])

                if q.get('chinese'):
                    p = doc.add_paragraph()
                    p.add_run(f"  中文: {q['chinese'][:200]}").font.color.rgb = RGBColor(0, 0, 255)

                for opt in ['A', 'B', 'C', 'D', 'E']:
                    if opt in q.get('options', {}):
                        p = doc.add_paragraph()
                        p.add_run(f"  {opt}. {q['options'][opt]['english'][:100]}")

                        if q.get('answer') == opt:
                            p.add_run(" ✓").bold =True
                            p.runs[-1].font.color.rgb = RGBColor(255, 0, 0)
                doc.add_paragraph('_' * 40)
    except Exception as e:
        print(f"Error {filename}: {e}")

# ================== Part 2: sqe_questions.js ==================
doc.add_page_break()
doc.add_heading('第二部分：sqe_questions.js 题库 (1000题)', level=1)

js_path = r'C:\Users\Administrator\Documents\sqe-test\sqe_questions.js'
with open(js_path, 'r', encoding='utf-8-sig') as f:
    content = f.read()
    # Handle both var and const declarations
    content = re.sub(r'^(?:const|var)\s+sqeQuestions\s*=\s*', '', content, flags=re.MULTILINE)
    content = content.strip().rstrip(';')

# Group by subject
subjects = {}
for q in questions:
    subj = q.get('category', 'Unknown')
    if subj not in subjects:
        subjects[subj] = []
    subjects[subj].append(q)

for subj, qs in sorted(subjects.items()):
    doc.add_heading(f'【{subj}】共{len(qs)}题', level=2)
    for q in qs[:100]:
        p = doc.add_paragraph()
        p.add_run(f"Q{q['number']}: ").bold = True
        p.add_run(q.get('stem', '')[:150] + '...')

        if q.get('stem_zh'):
            p = doc.add_paragraph()
            p.add_run(f"  中文: {q.get('stem_zh', '')[:150]}...").font.color.rgb = RGBColor(0, 0, 255)

        opts = q.get('opts', {})
        for opt in ['A', 'B', 'C', 'D', 'E']:
            if opt in opts:
                p = doc.add_paragraph()
                p.add_run(f"  {opt}. {opts[opt][:80]}...")
                if q.get('answer') == opt:
                    p.add_run(" ✓").bold = True
                    p.runs[-1].font.color.rgb = RGBColor(255, 0, 0)

        if q.get('analysis_zh'):
            p = doc.add_paragraph()
            p.add_run(f"  解析: {q['analysis_zh'][:100]}...").font.color.rgb = RGBColor(128, 0, 128)

        doc.add_paragraph('_' * 40)

# Save
output_path = r'C:\Users\Administrator\Desktop\SQE_Complete_Questions.docx'
doc.save(output_path)
print(f"\nDocument saved to: {output_path}")