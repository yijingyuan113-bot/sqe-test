import os
import re
import json
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

sqe_dir = r'C:\Users\Administrator\Desktop\Desktop\SQE'
docx_files = [f for f in os.listdir(sqe_dir) if f.endswith('.docx') and not f.startswith('words') and not f.startswith('千')]

def read_docx_full(path):
    doc = Document(path)
    texts = []
    for para in doc.paragraphs:
        if para.text.strip():
            texts.append(para.text)
    return '\n'.join(texts)

def parse_questions(text):
    """Parse questions from text - extract English, Chinese, options"""
    questions = []
    # Split by question numbers
    pattern = r'Question\s*(\d+)|^\d+\s*$'
    parts = re.split(pattern, text, flags=re.MULTILINE | re.IGNORECASE)

    current_q = None
    for part in parts:
        part = part.strip()
        if not part:
            continue
        if part.isdigit():
            if current_q:
                questions.append(current_q)
            current_q = {'number': int(part), 'english': '', 'chinese': '', 'options': {}}
        else:
            if current_q is None:
                current_q = {'number': 0, 'english': '', 'chinese': '', 'options': {}}
            current_q['raw'] = part
    if current_q and current_q.get('english'):
        questions.append(current_q)
    return questions

def parse_questions_v2(text):
    """Better question parser"""
    questions = []
    lines = text.split('\n')

    i = 0
    current_q = None
    section = 'unknown'

    while i < len(lines):
        line = lines[i].strip()

        # Question number patterns
        q_match = re.match(r'^Ques?tion\s*(\d+)', line, re.IGNORECASE)
        num_match = re.match(r'^(\d+)\s*$', line)

        if q_match or (num_match and not current_q):
            if current_q and current_q.get('english'):
                questions.append(current_q)
            num = int(q_match.group(1)) if q_match else int(num_match.group(1))
            current_q = {'number': num, 'english': '', 'chinese': '', 'options': {}, 'answer': ''}
            section = 'question'

        elif current_q:
            # Option detection
            opt_match = re.match(r'^([A-E])\.\s*(.+)', line, re.IGNORECASE)
            if opt_match:
                current_q['options'][opt_match.group(1).upper()] = {
                    'english': opt_match.group(2).strip(),
                    'chinese': ''
                }
                section = 'option'

            elif line.startswith('Answer') or line.startswith('正确答案'):
                section = 'answer'
                ans_match = re.search(r'([A-E])', line, re.IGNORECASE)
                if ans_match:
                    current_q['answer'] = ans_match.group(1).upper()

            elif 'english' in section or section == 'question':
                if line and not line.startswith('Options') and not line.startswith('Chinese'):
                    current_q['english'] += ' ' + line

            elif 'chinese' in section:
                if line:
                    current_q['chinese'] += ' ' + line

        i += 1

    if current_q and current_q.get('english'):
        questions.append(current_q)

    return questions

# Create combined Word document
doc = Document()
title = doc.add_heading('SQE 题目中英文对照及解析', 0)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# Process each docx file
subject_areas = {}
for filename in docx_files:
    filepath = os.path.join(sqe_dir, filename)
    print(f"Processing: {filename}")

    try:
        text = read_docx_full(filepath)
        questions = parse_questions_v2(text)
        print(f"  Found {len(questions)} questions")

        # Get subject area from filename
        subject = filename.replace('.docx', '').replace(' ', '_')
        subject_areas[subject] = questions

    except Exception as e:
        print(f"  Error: {e}")

print(f"\nTotal subject areas: {len(subject_areas)}")

# Add questions to Word doc by subject
for subject, questions in subject_areas.items():
    doc.add_heading(subject.replace('_', ' '), level=1)

    for q in questions[:50]:  # Limit to 50 per subject
        doc.add_heading(f"Q{q['number']}", level=2)

        # English
        p = doc.add_paragraph()
        p.add_run('English: ').bold = True
        p.add_run(q.get('english', ''))

        # Chinese
        chinese = q.get('chinese', '')
        if chinese:
            p = doc.add_paragraph()
            p.add_run('中文: ').bold = True
            p.add_run(chinese).font.color.rgb = RGBColor(0, 0, 255)

        # Options
        for opt_key in ['A', 'B', 'C', 'D', 'E']:
            if opt_key in q.get('options', {}):
                opt = q['options'][opt_key]
                p = doc.add_paragraph()
                p.add_run(f"{opt_key}. {opt['english']}").bold = True

                if opt.get('chinese'):
                    p = doc.add_paragraph()
                    p.add_run(f"   中文: {opt['chinese']}").font.color.rgb = RGBColor(0, 100, 0)

                if q.get('answer') == opt_key:
                    p = doc.add_paragraph()
                    p.add_run(f"   ✓ 正确答案").bold = True
                    p.runs[-1].font.color.rgb = RGBColor(255, 0, 0)

        doc.add_paragraph()

    doc.add_page_break()

# Save
output_path = r'C:\Users\Administrator\Desktop\SQE_题目汇总.docx'
doc.save(output_path)
print(f"\nDocument saved to: {output_path}")

# Save questions as JSON for web
all_questions = []
q_id = 1
for subject, questions in subject_areas.items():
    for q in questions:
        all_questions.append({
            'id': q_id,
            'number': q['number'],
            'subject': subject.replace('_', ' '),
            'flk': 'FLK1' if 'FLK1' in subject or 'FLK2' not in subject else 'FLK2',
            'stem': q.get('english', ''),
            'stem_zh': q.get('chinese', ''),
            'opts': {k: {'en': v['english'], 'zh': v.get('chinese', '')} for k, v in q.get('options', {}).items()},
            'answer': q.get('answer', '')
        })
        q_id += 1

json_path = r'C:\Users\Administrator\Desktop\SQE_questions.json'
with open(json_path, 'w', encoding='utf-8') as f:
    json.dump(all_questions, f, ensure_ascii=False, indent=2)
print(f"JSON saved to: {json_path}")
print(f"Total questions: {len(all_questions)}")